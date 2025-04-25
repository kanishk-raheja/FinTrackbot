from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_report():
    try:
        print("Starting to create the report...")
        # Create a new Document
        doc = Document()
        print("Document created successfully")
        
        # Add title
        title = doc.add_heading('GreenBuddy: A Specialized Gardening Chatbot - Technical Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        print("Title added")
        
        # Add introduction section
        doc.add_heading('1. Introduction', 1)
        intro_text = """GreenBuddy is a specialized chatbot designed to provide expert gardening advice and information to users. The application leverages the power of Google's Gemini AI model to deliver detailed, formatted responses to gardening-related questions. The chatbot is implemented as a web application with a Flask backend and a responsive HTML/CSS/JavaScript frontend, offering users multiple ways to interact with it - either through a floating widget on the main page or via a dedicated full-page experience.

The project aims to make gardening knowledge accessible to both beginners and experienced gardeners, providing practical advice on plant care, garden maintenance, and horticultural techniques. By focusing exclusively on gardening topics, GreenBuddy offers more specialized and relevant information compared to general-purpose chatbots."""
        doc.add_paragraph(intro_text)
        print("Introduction section added")
        
        # Add objective section
        doc.add_heading('2. Objective', 1)
        objective_text = """The primary objectives of the GreenBuddy chatbot are:

1. To provide accurate, practical, and helpful gardening advice to users of all experience levels
2. To create an intuitive and user-friendly interface for accessing gardening information
3. To leverage advanced AI capabilities to deliver detailed, contextually relevant responses
4. To offer multiple access options (floating widget and full-page experience) to accommodate different user preferences
5. To implement a responsive design that works well across various devices and screen sizes
6. To include accessibility features such as text-to-speech functionality and dark mode support"""
        doc.add_paragraph(objective_text)
        print("Objective section added")
        
        # Add data and API tuning section
        doc.add_heading('3. Data and API Tuning', 1)
        
        # Gemini API Token Limit
        doc.add_heading('Gemini API Token Limit', 2)
        token_text = """The GreenBuddy chatbot utilizes Google's Gemini 1.5 Pro model, which has the following token limits:
- Input token limit: 32,768 tokens
- Output token limit: 2,048 tokens
- Context window: 1 million tokens

The system prompt used to define the chatbot's personality and expertise is carefully crafted to fit within these limits while providing sufficient context for the model to generate appropriate responses."""
        doc.add_paragraph(token_text)
        
        # Data/Message Structure
        doc.add_heading('Data/Message Structure for the Chatbot', 2)
        structure_text = """The chatbot's message structure follows a simple JSON format:

Request Structure:
{
  "message": "User's gardening question or statement"
}

Response Structure:
{
  "response": "Bot's formatted gardening advice",
  "status": "success" | "error" | "rate_limit"
}

In case of errors, additional information is included:
{
  "response": "Error message for the user",
  "status": "error",
  "error": "Detailed error information for debugging"
}"""
        doc.add_paragraph(structure_text)
        
        # Storage Configuration
        doc.add_heading('Storage Configuration (Saving Chats)', 2)
        storage_text = """The current implementation does not include persistent storage for chat history. Each conversation is stateless, with the system prompt being sent with each new message to maintain context. This approach was chosen for simplicity and to avoid potential privacy concerns related to storing user conversations."""
        doc.add_paragraph(storage_text)
        
        # API Parameters
        doc.add_heading('API Parameters', 2)
        params_text = """The Gemini API is configured with the following parameters:

1. Model Selection: Gemini 1.5 Pro, which offers a good balance between performance and cost
2. System Prompt: A detailed prompt that defines the chatbot's expertise, personality, and response guidelines
3. Temperature: Default temperature setting (not explicitly specified in the code)
4. Max Output Tokens: Not explicitly limited, allowing the model to generate comprehensive responses"""
        doc.add_paragraph(params_text)
        
        # API Configuration
        doc.add_heading('API Configuration', 2)
        config_text = """The API configuration is handled through the following steps:

1. API Key Management: The Gemini API key is stored in a .env file for security and loaded using the python-dotenv library
2. Model Initialization: The model is initialized using genai.GenerativeModel('gemini-1.5-pro')
3. Conversation Management: Each user message creates a new chat session with the system prompt to ensure consistent context
4. Error Handling: Comprehensive error handling for API rate limits, network issues, and other potential problems"""
        doc.add_paragraph(config_text)
        print("Data and API Tuning section added")
        
        # Add methodology section
        doc.add_heading('4. Methodology', 1)
        
        # Backend Implementation
        doc.add_heading('Backend Implementation (Flask)', 2)
        backend_text = """1. Server Setup: A Flask server with CORS support to allow cross-origin requests
2. API Integration: Integration with the Gemini API using the Google Generative AI Python library
3. Request Processing: Processing of incoming chat messages and formatting of responses
4. Error Handling: Robust error handling for various failure scenarios"""
        doc.add_paragraph(backend_text)
        
        # Frontend Implementation
        doc.add_heading('Frontend Implementation', 2)
        frontend_text = """1. User Interface: A responsive UI with both light and dark mode support
2. Access Options: Two ways to access the chatbot:
   - Floating widget on the main page
   - Dedicated full-page experience
3. Interactive Elements: 
   - Text input field for user questions
   - Send button for submitting questions
   - Text-to-speech toggle for accessibility
   - Dark mode toggle for user preference
4. Visual Feedback: Typing indicators and error messages to enhance user experience"""
        doc.add_paragraph(frontend_text)
        
        # Communication Flow
        doc.add_heading('Communication Flow', 2)
        flow_text = """1. User enters a gardening question in the input field
2. The frontend sends a POST request to the Flask backend
3. The backend creates a new chat session with the Gemini model
4. The system prompt is sent to establish the chatbot's context
5. The user's question is sent to the model
6. The model generates a response based on the context and question
7. The response is sent back to the frontend
8. The frontend displays the response in the chat interface"""
        doc.add_paragraph(flow_text)
        print("Methodology section added")
        
        # Add results section
        doc.add_heading('6. Results', 1)
        results_text = """The GreenBuddy chatbot successfully achieves its objectives with the following results:

1. Functionality: The chatbot provides detailed, gardening-focused responses to user queries
2. User Experience: The interface is intuitive and responsive, with smooth transitions and visual feedback
3. Accessibility: Text-to-speech functionality and dark mode support enhance accessibility
4. Performance: The chatbot responds quickly to user queries, with appropriate loading indicators
5. Error Handling: The system gracefully handles API errors and rate limits
6. Flexibility: Users can choose between a floating widget and a full-page experience based on their preference

The chatbot has been tested with various gardening-related questions and consistently provides helpful, accurate information. The system prompt effectively guides the model to stay focused on gardening topics and provide practical advice."""
        doc.add_paragraph(results_text)
        print("Results section added")
        
        # Add conclusion section
        doc.add_heading('7. Conclusion', 1)
        conclusion_text = """GreenBuddy demonstrates the effective application of advanced AI technology to create a specialized, user-friendly chatbot for gardening advice. By leveraging the Gemini API and implementing a well-designed frontend, the project successfully delivers valuable gardening information to users in an accessible and engaging format.

The dual-access approach (floating widget and full-page experience) provides flexibility for different user preferences and use cases. The focus on gardening expertise ensures that users receive relevant, practical advice rather than generic responses.

Future enhancements could include:
1. Persistent chat history to maintain context across sessions
2. Image recognition capabilities to help users identify plants
3. Integration with gardening databases for more specific plant information
4. Personalized advice based on user location and climate
5. Community features to allow users to share gardening experiences

Overall, GreenBuddy represents a successful implementation of a specialized AI chatbot that provides genuine value to users interested in gardening."""
        doc.add_paragraph(conclusion_text)
        print("Conclusion section added")
        
        # Save the document
        output_path = 'GreenBuddy_Technical_Report.docx'
        doc.save(output_path)
        print(f"Report has been created as '{output_path}'")
        
        # Verify the file exists
        if os.path.exists(output_path):
            file_size = os.path.getsize(output_path)
            print(f"File exists with size: {file_size} bytes")
        else:
            print("Error: File was not created!")
            
    except Exception as e:
        print(f"Error creating report: {str(e)}")

if __name__ == "__main__":
    create_report() 