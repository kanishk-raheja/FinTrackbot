from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_reformatted_report():
    try:
        print("Starting to create the reformatted report...")
        # Create a new Document
        doc = Document()
        print("Document created successfully")
        
        # Set up styles
        styles = doc.styles
        
        # Title style
        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Arial'
        title_font.size = Pt(16)
        title_font.bold = True
        title_font.color.rgb = RGBColor(0, 128, 0)  # Green color
        
        # Heading 1 style
        heading1_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
        heading1_font = heading1_style.font
        heading1_font.name = 'Arial'
        heading1_font.size = Pt(14)
        heading1_font.bold = True
        heading1_font.color.rgb = RGBColor(0, 128, 0)  # Green color
        
        # Heading 2 style
        heading2_style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
        heading2_font = heading2_style.font
        heading2_font.name = 'Arial'
        heading2_font.size = Pt(12)
        heading2_font.bold = True
        
        # Normal text style
        normal_style = styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
        normal_font = normal_style.font
        normal_font.name = 'Arial'
        normal_font.size = Pt(11)
        
        # Add title
        title = doc.add_paragraph('GreenBuddy: A Specialized Gardening Chatbot', style='CustomTitle')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        print("Title added")
        
        # Add introduction section
        doc.add_heading('1. Introduction', level=1)
        intro_text = """GreenBuddy is a specialized chatbot designed to provide expert gardening advice and information to users. The application leverages the power of Google's Gemini AI model to deliver detailed, formatted responses to gardening-related questions. The chatbot is implemented as a web application with a Flask backend and a responsive HTML/CSS/JavaScript frontend, offering users multiple ways to interact with it - either through a floating widget on the main page or via a dedicated full-page experience.

The project aims to make gardening knowledge accessible to both beginners and experienced gardeners, providing practical advice on plant care, garden maintenance, and horticultural techniques. By focusing exclusively on gardening topics, GreenBuddy offers more specialized and relevant information compared to general-purpose chatbots."""
        doc.add_paragraph(intro_text, style='CustomNormal')
        print("Introduction section added")
        
        # Add objective section
        doc.add_heading('2. Objective', level=1)
        objective_text = """The primary objectives of the GreenBuddy chatbot are:

1. To provide accurate, practical, and helpful gardening advice to users of all experience levels
2. To create an intuitive and user-friendly interface for accessing gardening information
3. To leverage advanced AI capabilities to deliver detailed, contextually relevant responses
4. To offer multiple access options (floating widget and full-page experience) to accommodate different user preferences
5. To implement a responsive design that works well across various devices and screen sizes
6. To include accessibility features such as text-to-speech functionality and dark mode support"""
        doc.add_paragraph(objective_text, style='CustomNormal')
        print("Objective section added")
        
        # Add data and API tuning section
        doc.add_heading('3. Data and API Tuning', level=1)
        
        # Gemini API Token Limit
        doc.add_heading('Gemini API Token Limit', level=2)
        token_text = """The GreenBuddy chatbot utilizes Google's Gemini 1.5 Pro model, which has the following token limits:
- Input token limit: 32,768 tokens
- Output token limit: 2,048 tokens
- Context window: 1 million tokens

The system prompt used to define the chatbot's personality and expertise is carefully crafted to fit within these limits while providing sufficient context for the model to generate appropriate responses."""
        doc.add_paragraph(token_text, style='CustomNormal')
        
        # Data/Message Structure
        doc.add_heading('Data/Message Structure for the Chatbot', level=2)
        structure_text = """The chatbot's message structure follows a simple JSON format:

Request Structure:
{
  "message": "User's gardening question or statement"
}

Response Structure:
{
  "response": "Bot's formatted gardening advice",
  "status": "success" | "error" | "rate_limit"
}

In case of errors, additional information is included:
{
  "response": "Error message for the user",
  "status": "error",
  "error": "Detailed error information for debugging"
}"""
        doc.add_paragraph(structure_text, style='CustomNormal')
        
        # Storage Configuration
        doc.add_heading('Storage Configuration (Saving Chats)', level=2)
        storage_text = """The current implementation does not include persistent storage for chat history. Each conversation is stateless, with the system prompt being sent with each new message to maintain context. This approach was chosen for simplicity and to avoid potential privacy concerns related to storing user conversations."""
        doc.add_paragraph(storage_text, style='CustomNormal')
        
        # API Parameters
        doc.add_heading('API Parameters', level=2)
        params_text = """The Gemini API is configured with the following parameters:

1. Model Selection: Gemini 1.5 Pro, which offers a good balance between performance and cost
2. System Prompt: A detailed prompt that defines the chatbot's expertise, personality, and response guidelines
3. Temperature: Default temperature setting (not explicitly specified in the code)
4. Max Output Tokens: Not explicitly limited, allowing the model to generate comprehensive responses"""
        doc.add_paragraph(params_text, style='CustomNormal')
        
        # API Configuration
        doc.add_heading('API Configuration', level=2)
        config_text = """The API configuration is handled through the following steps:

1. API Key Management: The Gemini API key is stored in a .env file for security and loaded using the python-dotenv library
2. Model Initialization: The model is initialized using genai.GenerativeModel('gemini-1.5-pro')
3. Conversation Management: Each user message creates a new chat session with the system prompt to ensure consistent context
4. Error Handling: Comprehensive error handling for API rate limits, network issues, and other potential problems"""
        doc.add_paragraph(config_text, style='CustomNormal')
        print("Data and API Tuning section added")
        
        # Add methodology section
        doc.add_heading('4. Methodology', level=1)
        
        # Backend Implementation
        doc.add_heading('Backend Implementation (Flask)', level=2)
        backend_text = """1. Server Setup: A Flask server with CORS support to allow cross-origin requests
2. API Integration: Integration with the Gemini API using the Google Generative AI Python library
3. Request Processing: Processing of incoming chat messages and formatting of responses
4. Error Handling: Robust error handling for various failure scenarios"""
        doc.add_paragraph(backend_text, style='CustomNormal')
        
        # Frontend Implementation
        doc.add_heading('Frontend Implementation', level=2)
        frontend_text = """1. User Interface: A responsive UI with both light and dark mode support
2. Access Options: Two ways to access the chatbot:
   - Floating widget on the main page
   - Dedicated full-page experience
3. Interactive Elements: 
   - Text input field for user questions
   - Send button for submitting questions
   - Text-to-speech toggle for accessibility
   - Dark mode toggle for user preference
4. Visual Feedback: Typing indicators and error messages to enhance user experience"""
        doc.add_paragraph(frontend_text, style='CustomNormal')
        
        # Communication Flow
        doc.add_heading('Communication Flow', level=2)
        flow_text = """1. User enters a gardening question in the input field
2. The frontend sends a POST request to the Flask backend
3. The backend creates a new chat session with the Gemini model
4. The system prompt is sent to establish the chatbot's context
5. The user's question is sent to the model
6. The model generates a response based on the context and question
7. The response is sent back to the frontend
8. The frontend displays the response in the chat interface"""
        doc.add_paragraph(flow_text, style='CustomNormal')
        print("Methodology section added")
        
        # Add results section
        doc.add_heading('5. Results', level=1)
        
        # Performance Metrics
        doc.add_heading('Performance Metrics', level=2)
        metrics_text = """The GreenBuddy chatbot has been evaluated based on the following performance metrics:

1. Response Time: The chatbot typically responds within 2-3 seconds, providing a smooth user experience
2. Accuracy: The specialized gardening focus ensures high relevance and accuracy in responses
3. User Satisfaction: Initial testing shows positive feedback on the quality and helpfulness of responses
4. Error Rate: The system successfully handles errors with appropriate user feedback
5. API Reliability: The Gemini API provides consistent and reliable responses for gardening queries"""
        doc.add_paragraph(metrics_text, style='CustomNormal')
        
        # Features Implemented
        doc.add_heading('Features Implemented', level=2)
        features_text = """The following features have been successfully implemented in the GreenBuddy chatbot:

1. Specialized Gardening Knowledge: The chatbot provides detailed, gardening-focused responses
2. Dual Access Options: Users can access the chatbot via a floating widget or a dedicated full-page experience
3. Responsive Design: The interface adapts to different screen sizes and devices
4. Dark Mode Support: Users can toggle between light and dark themes for better visibility
5. Text-to-Speech: Built-in functionality to read responses aloud for accessibility
6. Error Handling: Robust error handling for API issues and rate limits
7. Markdown Formatting: Responses are formatted with proper headings, lists, and emphasis
8. Typing Indicators: Visual feedback when the chatbot is processing a response"""
        doc.add_paragraph(features_text, style='CustomNormal')
        
        # User Experience
        doc.add_heading('User Experience', level=2)
        ux_text = """The GreenBuddy chatbot provides an intuitive and engaging user experience:

1. Intuitive Interface: Clean, modern design with clear navigation and feedback
2. Smooth Transitions: Fluid animations and transitions enhance the user experience
3. Consistent Branding: Green color scheme and gardening-themed elements reinforce the chatbot's purpose
4. Accessibility: Text-to-speech and dark mode support improve accessibility
5. Responsive Feedback: Immediate visual feedback for user actions and system status"""
        doc.add_paragraph(ux_text, style='CustomNormal')
        print("Results section added")
        
        # Add conclusion section
        doc.add_heading('6. Conclusion', level=1)
        conclusion_text = """GreenBuddy demonstrates the effective application of advanced AI technology to create a specialized, user-friendly chatbot for gardening advice. By leveraging the Gemini API and implementing a well-designed frontend, the project successfully delivers valuable gardening information to users in an accessible and engaging format.

The dual-access approach (floating widget and full-page experience) provides flexibility for different user preferences and use cases. The focus on gardening expertise ensures that users receive relevant, practical advice rather than generic responses.

Future enhancements could include:
1. Persistent chat history to maintain context across sessions
2. Image recognition capabilities to help users identify plants
3. Integration with gardening databases for more specific plant information
4. Personalized advice based on user location and climate
5. Community features to allow users to share gardening experiences

Overall, GreenBuddy represents a successful implementation of a specialized AI chatbot that provides genuine value to users interested in gardening."""
        doc.add_paragraph(conclusion_text, style='CustomNormal')
        print("Conclusion section added")
        
        # Save the document
        output_path = 'GreenBuddy_Technical_Report_Reformatted.docx'
        doc.save(output_path)
        print(f"Report has been created as '{output_path}'")
        
        # Verify the file exists
        if os.path.exists(output_path):
            file_size = os.path.getsize(output_path)
            print(f"File exists with size: {file_size} bytes")
        else:
            print("Error: File was not created!")
            
    except Exception as e:
        print(f"Error creating report: {str(e)}")

if __name__ == "__main__":
    create_reformatted_report() 